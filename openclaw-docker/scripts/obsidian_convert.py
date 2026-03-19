#!/usr/bin/env python3
"""
Obsidian to Markdown converter for OpenViking
Converts .docx and .pdf files to .md format
"""

import os
import sys
import json
from pathlib import Path
from datetime import datetime

# Try to import converters
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

def convert_docx(docx_path: Path, output_path: Path) -> bool:
    """Convert DOCX to Markdown"""
    if not HAS_DOCX:
        return False
    
    try:
        doc = Document(docx_path)
        md_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Handle headings
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    md_content.append(f"{'#' * int(level)} {text}")
                else:
                    md_content.append(text)
        
        # Handle tables
        for table in doc.tables:
            md_content.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                md_content.append("| " + " | ".join(cells) + " |")
        
        output_path.write_text("\n\n".join(md_content), encoding='utf-8')
        return True
    except Exception as e:
        print(f"Error converting {docx_path}: {e}", file=sys.stderr)
        return False

def convert_pdf(pdf_path: Path, output_path: Path) -> bool:
    """Convert PDF to Markdown"""
    if not HAS_PDF:
        return False
    
    try:
        md_content = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    md_content.append(f"## Page {page_num}\n\n{text}")
        
        output_path.write_text("\n\n".join(md_content), encoding='utf-8')
        return True
    except Exception as e:
        print(f"Error converting {pdf_path}: {e}", file=sys.stderr)
        return False

def find_and_convert(source_dir: str, output_dir: str, dry_run: bool = False):
    """Find and convert all supported files"""
    source_path = Path(source_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # File extensions to convert
    extensions = {'.docx': convert_docx, '.pdf': convert_pdf}
    
    converted = []
    skipped = []
    errors = []
    
    for ext, converter in extensions.items():
        for file_path in source_path.rglob(f'*{ext}'):
            # Create relative path structure
            rel_path = file_path.relative_to(source_path)
            md_path = output_path / rel_path.with_suffix('.md')
            
            # Create parent directories
            md_path.parent.mkdir(parents=True, exist_ok=True)
            
            if dry_run:
                skipped.append(str(rel_path))
                continue
            
            # Convert
            if converter(file_path, md_path):
                converted.append(str(rel_path))
            else:
                errors.append(str(rel_path))
    
    return {
        'converted': converted,
        'skipped': skipped,
        'errors': errors,
        'total': len(converted) + len(skipped) + len(errors)
    }

if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert Obsidian files to Markdown')
    parser.add_argument('--source', default='/data/obsidian', help='Source directory')
    parser.add_argument('--output', default='/data/obsidian_converted', help='Output directory')
    parser.add_argument('--dry-run', action='store_true', help='Show what would be converted')
    parser.add_argument('--json', action='store_true', help='Output JSON')
    
    args = parser.parse_args()
    
    result = find_and_convert(args.source, args.output, args.dry_run)
    
    if args.json:
        print(json.dumps(result, indent=2))
    else:
        print(f"Converted: {len(result['converted'])}")
        print(f"Skipped: {len(result['skipped'])}")
        print(f"Errors: {len(result['errors'])}")
        
        if result['converted']:
            print("\nRecently converted:")
            for f in result['converted'][-5:]:
                print(f"  + {f}")
